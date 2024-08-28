import os
import sys
import csv
import sqlite3
import numpy as np
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QMainWindow, QApplication, QPushButton, QLabel
from matplotlib import pyplot
from docx import Document

SPENDING = 0
EARNING = 1


class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(650, 390)
        MainWindow.setMinimumSize(QtCore.QSize(650, 390))
        font = QtGui.QFont()
        font.setFamily("Cambria")
        font.setPointSize(12)
        MainWindow.setFont(font)
        MainWindow.setTabShape(QtWidgets.QTabWidget.Rounded)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.centralwidget)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.verticalLayout = QtWidgets.QVBoxLayout()
        self.verticalLayout.setObjectName("verticalLayout")
        self.calendarWidget = QtWidgets.QCalendarWidget(self.centralwidget)
        font = QtGui.QFont()
        font.setFamily("Monotype Corsiva")
        font.setPointSize(15)
        font.setItalic(True)
        self.calendarWidget.setFont(font)
        self.calendarWidget.setLocale(QtCore.QLocale(QtCore.QLocale.Russian, QtCore.QLocale.Russia))
        self.calendarWidget.setMinimumDate(QtCore.QDate(2019, 1, 1))
        self.calendarWidget.setMaximumDate(QtCore.QDate(2199, 12, 31))
        self.calendarWidget.setHorizontalHeaderFormat(QtWidgets.QCalendarWidget.ShortDayNames)
        self.calendarWidget.setVerticalHeaderFormat(QtWidgets.QCalendarWidget.NoVerticalHeader)
        self.calendarWidget.setNavigationBarVisible(True)
        self.calendarWidget.setDateEditEnabled(True)
        self.calendarWidget.setObjectName("calendarWidget")
        self.verticalLayout.addWidget(self.calendarWidget)
        self.pushButton = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton.setObjectName("pushButton")
        self.verticalLayout.addWidget(self.pushButton)
        self.pushButton_2 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_2.setObjectName("pushButton_2")
        self.verticalLayout.addWidget(self.pushButton_2)
        self.horizontalLayout.addLayout(self.verticalLayout)
        self.listWidget = QtWidgets.QListWidget(self.centralwidget)
        font = QtGui.QFont()
        font.setFamily("Cambria")
        font.setPointSize(15)
        self.listWidget.setFont(font)
        self.listWidget.setObjectName("listWidget")
        self.horizontalLayout.addWidget(self.listWidget)
        MainWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MoneyKeeper"))
        self.pushButton.setText(_translate("MainWindow", "Добавить значение"))
        self.pushButton_2.setText(_translate("MainWindow", "Посмотреть аналитику"))


class Ui_DialogWindow(object):
    def setupUi(self, DialogWindow):
        DialogWindow.setObjectName("DialogWindow")
        DialogWindow.resize(250, 270)
        DialogWindow.setMaximumSize(QtCore.QSize(250, 270))
        DialogWindow.setMinimumSize(QtCore.QSize(250, 270))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(DialogWindow.sizePolicy().hasHeightForWidth())
        DialogWindow.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setFamily("Cambria")
        font.setPointSize(12)
        DialogWindow.setFont(font)
        self.centralwidget = QtWidgets.QWidget(DialogWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.comboBox = QtWidgets.QComboBox(self.centralwidget)
        self.comboBox.setObjectName("comboBox")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.verticalLayout.addWidget(self.comboBox)
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setAlignment(QtCore.Qt.AlignCenter)
        self.label.setObjectName("label")
        self.verticalLayout.addWidget(self.label)
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.lineEdit = QtWidgets.QLineEdit(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.lineEdit.sizePolicy().hasHeightForWidth())
        self.lineEdit.setSizePolicy(sizePolicy)
        self.lineEdit.setMaximumSize(QtCore.QSize(180, 16777215))
        self.lineEdit.setInputMask("")
        self.lineEdit.setText("")
        self.lineEdit.setAlignment(QtCore.Qt.AlignCenter)
        self.lineEdit.setObjectName("lineEdit")
        self.horizontalLayout.addWidget(self.lineEdit)
        self.verticalLayout.addLayout(self.horizontalLayout)
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setAlignment(QtCore.Qt.AlignCenter)
        self.label_2.setObjectName("label_2")
        self.verticalLayout.addWidget(self.label_2)
        self.horizontalLayout_2 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_2.setObjectName("horizontalLayout_2")
        self.lineEdit_2 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_2.setMaximumSize(QtCore.QSize(180, 16777215))
        self.lineEdit_2.setObjectName("lineEdit_2")
        self.horizontalLayout_2.addWidget(self.lineEdit_2)
        self.verticalLayout.addLayout(self.horizontalLayout_2)
        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.verticalLayout.addItem(spacerItem)
        self.pushButton = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton.setMinimumSize(QtCore.QSize(0, 35))
        self.pushButton.setObjectName("pushButton")
        self.verticalLayout.addWidget(self.pushButton)
        self.status = QtWidgets.QLabel(self.centralwidget)
        font = QtGui.QFont()
        font.setFamily("Cambria")
        font.setPointSize(10)
        self.status.setFont(font)
        self.status.setText("")
        self.status.setAlignment(QtCore.Qt.AlignCenter)
        self.status.setObjectName("label_3")
        self.verticalLayout.addWidget(self.status)
        DialogWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(DialogWindow)
        QtCore.QMetaObject.connectSlotsByName(DialogWindow)

    def retranslateUi(self, DialogWindow):
        _translate = QtCore.QCoreApplication.translate
        DialogWindow.setWindowTitle(_translate("DialogWindow", "Добавление значения"))
        self.comboBox.setItemText(0, _translate("DialogWindow", "Трата"))
        self.comboBox.setItemText(1, _translate("DialogWindow", "Заработок"))
        self.label.setText(_translate("DialogWindow", "Название"))
        self.label_2.setText(_translate("DialogWindow", "Количество, руб."))
        self.pushButton.setText(_translate("DialogWindow", "Добавить"))


class Ui_AnalyticsWindow(object):
    def setupUi(self, AnalyticsWindow):
        AnalyticsWindow.setObjectName("AnalyticsWindow")
        AnalyticsWindow.resize(250, 150)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(AnalyticsWindow.sizePolicy().hasHeightForWidth())
        AnalyticsWindow.setSizePolicy(sizePolicy)
        AnalyticsWindow.setMinimumSize(QtCore.QSize(250, 150))
        AnalyticsWindow.setMaximumSize(QtCore.QSize(250, 150))
        font = QtGui.QFont()
        font.setFamily("Cambria")
        font.setPointSize(12)
        AnalyticsWindow.setFont(font)
        self.centralwidget = QtWidgets.QWidget(AnalyticsWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout_4 = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout_4.setObjectName("verticalLayout_4")
        self.horizontalLayout_4 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_4.setObjectName("horizontalLayout_4")
        self.dateEdit = QtWidgets.QDateEdit(self.centralwidget)
        self.dateEdit.setDateTime(QtCore.QDateTime(QtCore.QDate(2021, 11, 27), QtCore.QTime(0, 0, 0)))
        self.dateEdit.setMaximumDate(QtCore.QDate(2199, 12, 31))
        self.dateEdit.setMinimumDate(QtCore.QDate(2019, 1, 1))
        self.dateEdit.setObjectName("dateEdit")
        self.horizontalLayout_4.addWidget(self.dateEdit)
        self.dateEdit_2 = QtWidgets.QDateEdit(self.centralwidget)
        self.dateEdit_2.setDateTime(QtCore.QDateTime(QtCore.QDate(2021, 11, 27), QtCore.QTime(0, 0, 0)))
        self.dateEdit_2.setMaximumDate(QtCore.QDate(2199, 12, 31))
        self.dateEdit_2.setMinimumDate(QtCore.QDate(2019, 1, 1))
        self.dateEdit_2.setObjectName("dateEdit_2")
        self.horizontalLayout_4.addWidget(self.dateEdit_2)
        self.verticalLayout_4.addLayout(self.horizontalLayout_4)
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.pushButton = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton.setMaximumSize(QtCore.QSize(170, 16777215))
        self.pushButton.setObjectName("pushButton")
        self.horizontalLayout.addWidget(self.pushButton)
        self.verticalLayout_4.addLayout(self.horizontalLayout)
        self.horizontalLayout_2 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_2.setObjectName("horizontalLayout_2")
        self.pushButton_2 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_2.setMaximumSize(QtCore.QSize(170, 16777215))
        self.pushButton_2.setObjectName("pushButton_2")
        self.horizontalLayout_2.addWidget(self.pushButton_2)
        self.verticalLayout_4.addLayout(self.horizontalLayout_2)
        self.horizontalLayout_3 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_3.setObjectName("horizontalLayout_3")
        self.pushButton_3 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_3.setMaximumSize(QtCore.QSize(170, 16777215))
        self.pushButton_3.setObjectName("pushButton_3")
        self.horizontalLayout_3.addWidget(self.pushButton_3)
        self.verticalLayout_4.addLayout(self.horizontalLayout_3)
        AnalyticsWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(AnalyticsWindow)
        QtCore.QMetaObject.connectSlotsByName(AnalyticsWindow)

    def retranslateUi(self, AnalyticsWindow):
        _translate = QtCore.QCoreApplication.translate
        AnalyticsWindow.setWindowTitle(_translate("AnalyticsWindow", "Посмотреть аналитику"))
        self.pushButton.setText(_translate("AnalyticsWindow", "Экспорт в word"))
        self.pushButton_2.setText(_translate("AnalyticsWindow", "Экспорт в csv"))
        self.pushButton_3.setText(_translate("AnalyticsWindow", "Экспорт в sql"))


class Ui_DeleteWindow(object):
    def setupUi(self, DeleteWindow):
        DeleteWindow.setObjectName("DeleteWindow")
        DeleteWindow.resize(280, 260)
        DeleteWindow.setMinimumSize(QtCore.QSize(280, 260))
        DeleteWindow.setMaximumSize(QtCore.QSize(280, 260))
        font = QtGui.QFont()
        font.setFamily("Cambria")
        font.setPointSize(14)
        DeleteWindow.setFont(font)
        self.centralwidget = QtWidgets.QWidget(DeleteWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.verticalLayout.addItem(spacerItem)
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setAlignment(QtCore.Qt.AlignCenter)
        self.label.setObjectName("label")
        self.verticalLayout.addWidget(self.label)
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setAlignment(QtCore.Qt.AlignCenter)
        self.label_2.setObjectName("label_2")
        self.verticalLayout.addWidget(self.label_2)
        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.verticalLayout.addItem(spacerItem1)
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.btn_yes = QtWidgets.QPushButton(self.centralwidget)
        self.btn_yes.setObjectName("pushButton_2")
        self.horizontalLayout.addWidget(self.btn_yes)
        self.btn_no = QtWidgets.QPushButton(self.centralwidget)
        self.btn_no.setObjectName("pushButton")
        self.horizontalLayout.addWidget(self.btn_no)
        self.verticalLayout.addLayout(self.horizontalLayout)
        DeleteWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(DeleteWindow)
        QtCore.QMetaObject.connectSlotsByName(DeleteWindow)

    def retranslateUi(self, DeleteWindow):
        _translate = QtCore.QCoreApplication.translate
        DeleteWindow.setWindowTitle(_translate("DeleteWindow", "MainWindow"))
        self.label.setText(_translate("DeleteWindow", "Вы уверены, что хотите"))
        self.label_2.setText(_translate("DeleteWindow", "удалить данный элемент?"))
        self.btn_yes.setText(_translate("DeleteWindow", "Да"))
        self.btn_no.setText(_translate("DeleteWindow", "Нет"))


class MyWidget(QMainWindow, Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.pushButton.clicked.connect(self.onClick)
        self.pushButton_2.clicked.connect(self.view_analytics)
        self.calendarWidget.selectionChanged.connect(self.onDateChanged)
        self.listWidget.clicked.connect(self.onListItemChanged)
        self.data = {}
        self.date = ""
        self.load_from_logs()

    def load_from_logs(self):
        file = open("logs.csv", encoding="utf-8")
        reader = csv.reader(file, delimiter=";", quotechar='\"')
        for i in reader:
            if i[0] not in self.data.keys():
                self.data[i[0]] = [": ".join(i[1:])]
            else:
                self.data[i[0]].append(": ".join(i[1:]))
        file.close()
        self.onDateChanged()

    def onDateChanged(self):
        date = self.calendarWidget.selectedDate()
        date = date.toString("dd.MM.yyyy")

        # if self.date != date:
        self.date = date
        if date not in self.data.keys():
            self.data[date] = []
        self.listWidget.clear()
        self.listWidget.addItems(self.data[date])

    def onListItemChanged(self):
        # Инициализация диалогового окна

        self.deleting_item = self.listWidget.currentItem().text()
        self.dialog = DeleteDialog(self)
        self.dialog.show()

    def init_dialog(self):
        # Инициализация диалогового окна

        self.dialog = MyDialog(self)
        self.dialog.show()

    def init_analytics(self):
        # Инициализация окна просмотра аналитики

        self.dialog = Analytics(self.data)
        self.dialog.show()

    def load_res(self, type, name, value):
        # Загрузка результатов из диалога добавления

        if type == SPENDING:
            item = f"- {value}: {name}"
        elif type == EARNING:
            item = f"+ {value}: {name}"
        self.data[self.date].append(item)

        self.listWidget.clear()
        self.listWidget.addItems(self.data[self.date])

    def delete(self, b):
        if b:
            k = self.data[self.date].index(self.deleting_item)
            del self.data[self.date][k]
            self.onDateChanged()
        self.deleting_item = None


    def onClick(self):
        self.onDateChanged()
        self.init_dialog()

    def view_analytics(self):
        self.onDateChanged()
        self.init_analytics()

    def closeEvent(self, a0: QtGui.QCloseEvent) -> None:
        file = open("logs.csv", "w", encoding="utf-8", newline="")
        writer = csv.writer(file, delimiter=";", quotechar='\"')

        for key, value in self.data.items():
            for j in value:
                k = [key] + j.split(": ")
                writer.writerow(k)
        file.close()


class MyDialog(QMainWindow, Ui_DialogWindow):
    def __init__(self, parent):
        super().__init__()

        self.setupUi(self)
        self.pushButton.clicked.connect(self.onClick)
        self.parent = parent

    def onClick(self):
        name = self.lineEdit.text()
        value = self.lineEdit_2.text()
        type = SPENDING if self.comboBox.currentText() == "Трата" else EARNING

        if name and value and value.isdigit():
            self.close()
            self.parent.load_res(type, name, value)
        else:
            self.status.setText("Заполните все поля формы правильно")


class Analytics(QMainWindow, Ui_AnalyticsWindow):
    def __init__(self, data):
        super().__init__()

        self.setupUi(self)
        self.pushButton.clicked.connect(self.word)
        self.pushButton_2.clicked.connect(self.csv)
        self.pushButton_3.clicked.connect(self.sql)
        self.data = data

    def convert_date(self, d1):
        d1 = d1.split(".")
        d1 = int(d1[0]) + int(d1[1]) * 30 + (int(d1[2]) - 2019) * 365
        return d1

    def prepare_to_convert(self):
        # Функция генерирует словарь, который обрабатывается
        # функциями промотра аналитики
        d1 = self.convert_date(self.dateEdit.text())
        d2 = self.convert_date(self.dateEdit_2.text())
        date_range = range(d1, d2 + 1)

        lst = self.data.items()
        lst = list(filter(lambda x: self.convert_date(x[0]) in date_range, lst))
        d = {}
        for i in lst:
            for j in i[1]:
                k = int("".join(j.split(": ")[0].split(" ")))
                if i[0] in d.keys():
                    if k > 0:
                        d[i[0]][0] += k
                    else:
                        d[i[0]][1] += -k
                else:
                    d[i[0]] = [0, 0]
                    if k > 0:
                        d[i[0]][0] += k
                    else:
                        d[i[0]][1] += -k

        lst = list(d.items())
        lst.sort(key=lambda x: x[0])
        d = {}

        for i in lst:
            d[i[0]] = i[1]
        return d

    def word(self):
        d = self.prepare_to_convert()

        labels = d.keys()
        earnings = [i[0] for i in d.values()]
        spendings = [i[1] for i in d.values()]

        x = np.arange(len(labels))
        width = 0.35

        fig, ax = pyplot.subplots()

        rects1 = ax.bar(x - width / 2, earnings, width, label="Заработки")
        rects2 = ax.bar(x + width / 2, spendings, width, label="Траты")

        ax.set_ylabel("")
        ax.set_title("")
        ax.set_xticks(x, labels)
        ax.legend()

        ax.bar_label(rects1, padding=3)
        ax.bar_label(rects2, padding=3)

        fig.tight_layout()

        pyplot.savefig("analytics.png")

        document = Document()
        paragraph = document.add_paragraph()

        run = paragraph.add_run()
        run.add_text("Аналитика")
        run.add_picture("analytics.png")

        document.save("analytics.docx")
        os.remove("analytics.png")

        self.close()

    def csv(self):
        d = self.prepare_to_convert()

        file = open("analytics.csv", "w", encoding="utf-8", newline="")
        writer = csv.writer(file, delimiter=";", quotechar='\"')

        writer.writerow(["дата", "прибыль", "убытки"])
        for key, value in d.items():
            writer.writerow([key, value[0], value[1]])
        file.close()

        self.close()

    def sql(self):
        d = self.prepare_to_convert()

        file = open("analytics.db", "w")
        file.close()

        connection = sqlite3.connect("analytics.db")
        cursor = connection.cursor()
        cursor.execute("CREATE TABLE Analytics (Date TEXT, Earning INT, Spending INT)")

        for key, value in d.items():
            cursor.execute(f"INSERT INTO Analytics(Date, Earning, Spending) VALUES('{key}', {value[0]}, {value[1]})")
        connection.commit()
        connection.close()

        self.close()


class DeleteDialog(QMainWindow, Ui_DeleteWindow):
    def __init__(self, parent):
        super().__init__()

        self.setupUi(self)

        self.btn_no.clicked.connect(self.onClick)
        self.btn_yes.clicked.connect(self.onClick)

        self.parent = parent

    def onClick(self):
        if self.sender().text() == "Да":
            self.parent.delete(True)
        else:
            self.parent.delete(False)
        self.close()


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = MyWidget()
    ex.show()
    sys.exit(app.exec_())
